import requests
import time
import os
import re
import tempfile
import zipfile
from urllib.parse import urljoin, urlparse
from bs4 import BeautifulSoup, NavigableString
from docx import Document as WordDocument
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches

HEADERS = {
    "User-Agent": "Mozilla/5.0"
}

# -----------------------------
# FETCH HTML
# -----------------------------
def fetch_html(url):
    response = requests.get(url, headers=HEADERS)
    if response.status_code != 200:
        print(f"⚠️ Failed to fetch {url} (Status code: {response.status_code})")
    response.raise_for_status()
    return response.text


# -----------------------------
# TITLE
# -----------------------------
def extract_title(soup):
    if soup.title and soup.title.string:
        return soup.title.string.strip().split("|")[0].strip()
    return "Scraped_Page"


def sanitize_filename(title):
    title = re.sub(r'[\\/*?:"<>|]', "", title)
    return title.replace(" ", "_")[:100]


# -----------------------------
# FOLDER SETUP
# -----------------------------
def create_page_folder(title, base_dir):
    folder_name = sanitize_filename(title)
    folder_path = os.path.join(base_dir, folder_name)
    os.makedirs(folder_path, exist_ok=True)

    images_folder = os.path.join(folder_path, "images")
    os.makedirs(images_folder, exist_ok=True)

    return folder_path, images_folder


# -----------------------------
# CLEAN TEXT
# -----------------------------
def clean_text(text):
    return " ".join(text.split())

# -----------------------------
# 🎥 YOUTUBE FUNCTIONS (KEPT)
# -----------------------------
def extract_youtube_id(url):
    patterns = [
        r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([^&\n?#]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def extract_youtube_videos(soup):
    videos = []

    # From iframes
    for iframe in soup.find_all("iframe"):
        src = iframe.get("src", "")
        video_id = extract_youtube_id(src)
        if video_id:
            videos.append({
                "id": video_id,
                "url": f"https://www.youtube.com/watch?v={video_id}",
                "title": iframe.get("title", "YouTube Video")
            })

    # From links
    for link in soup.find_all("a", href=True):
        href = link.get("href", "")
        if "youtube.com" in href or "youtu.be" in href:
            video_id = extract_youtube_id(href)
            if video_id:
                videos.append({
                    "id": video_id,
                    "url": f"https://www.youtube.com/watch?v={video_id}",
                    "title": link.get_text(strip=True) or "YouTube Video"
                })

    return videos



# -----------------------------
# DOWNLOAD IMAGES
# -----------------------------
def download_images(soup, base_url, images_folder):
    image_map = {}
    images = soup.find_all("img")

    for idx, img in enumerate(images):
        src = img.get("src") or img.get("data-src")
        if not src:
            continue

        img_url = urljoin(base_url, src)

        try:
            response = requests.get(img_url, headers=HEADERS, timeout=10)
            response.raise_for_status()

            parsed = urlparse(img_url)
            ext = os.path.splitext(parsed.path)[1] or ".jpg" or ".jpeg" or ".png" or ".svg"

            # retain the image file name
            filename = os.path.basename(parsed.path)
            if not filename:
                filename = f"image_{idx}{ext}"
            filepath = os.path.join(images_folder, filename)

            with open(filepath, "wb") as f:
                f.write(response.content)

            image_map[src] = filepath

        except Exception as e:
            print(f"⚠️ Failed to download image: {img_url} ({e})")

    return image_map


# -----------------------------
# CONTENT EXTRACTION
# -----------------------------
def extract_main_content(soup):
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        tag.decompose()

    return (
        soup.find("main") or
        soup.find("article") or
        soup.find("div", {"id": "__next"}) or
        soup.body
    )


# -----------------------------
# HYPERLINK HANDLER
# -----------------------------
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


# -----------------------------
# INLINE CONTENT
# -----------------------------
def add_inline_content(paragraph, element):
    for node in element.children:

        if isinstance(node, NavigableString):
            text = clean_text(str(node))
            if text:
                paragraph.add_run(text + " ")

        elif node.name == "a":
            text = clean_text(node.get_text())
            href = node.get("href")
            if text and href:
                add_hyperlink(paragraph, text, href)
                paragraph.add_run(" ")

        elif node.name in ["strong", "b"]:
            text = clean_text(node.get_text())
            if text:
                run = paragraph.add_run(text + " ")
                run.bold = True

        elif node.name in ["em", "i"]:
            text = clean_text(node.get_text())
            if text:
                run = paragraph.add_run(text + " ")
                run.italic = True

        else:
            add_inline_content(paragraph, node)


# -----------------------------
# TABLE HANDLING
# -----------------------------
def add_table_to_doc(doc, table_element):
    rows = table_element.find_all("tr")
    if not rows:
        return

    max_cols = 0
    parsed_rows = []

    for row in rows:
        cells = row.find_all(["th", "td"])
        parsed_cells = []

        for cell in cells:
            colspan = int(cell.get("colspan", 1))
            parsed_cells.append((cell, colspan))

        parsed_rows.append(parsed_cells)
        max_cols = max(max_cols, sum(c for _, c in parsed_cells))

    table = doc.add_table(rows=0, cols=max_cols)
    table.style = "Table Grid"

    for row_data in parsed_rows:
        row_cells = table.add_row().cells
        col_index = 0

        for cell_element, colspan in row_data:
            cell = row_cells[col_index]
            paragraph = cell.paragraphs[0]

            add_inline_content(paragraph, cell_element)

            if colspan > 1:
                for i in range(1, colspan):
                    cell.merge(row_cells[col_index + i])

            col_index += colspan

    doc.add_paragraph("")



# -----------------------------
# BUILD DOCUMENT
# -----------------------------
def build_document(content, title, url, filename, image_map):
    doc = WordDocument()

    doc.add_heading(title, 0)
    doc.add_paragraph(f"Source: {url}")

    # 🎥 Add YouTube videos at top
    videos = extract_youtube_videos(content)
    if videos:
        doc.add_heading("YouTube Videos", level=2)
        for video in videos:
            p = doc.add_paragraph()
            p.add_run("🎥 ")
            add_hyperlink(p, video["title"], video["url"])
            p.add_run(f" ({video['id']})")
        doc.add_paragraph("")

    for element in content.find_all(
        ["h1", "h2", "h3", "h4", "p", "ul", "ol", "table", "img"]
    ):

        if element.name in ["h1", "h2", "h3", "h4"]:
            text = clean_text(element.get_text())
            if text:
                doc.add_heading(text, level=int(element.name[1]))

        elif element.name == "p":
            p = doc.add_paragraph()
            add_inline_content(p, element)

        elif element.name in ["ul", "ol"]:
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(
                    style="List Bullet" if element.name == "ul" else "List Number"
                )
                add_inline_content(p, li)

        elif element.name == "table":
            add_table_to_doc(doc, element)

        elif element.name == "img":
            src = element.get("src") or element.get("data-src")
            if src and src in image_map:
                try:
                    doc.add_picture(image_map[src], width=Inches(5))
                except Exception as e:
                    print(f"⚠️ Failed to embed image: {src} ({e})")

    doc.save(filename)


# -----------------------------
# MAIN PROCESS
# -----------------------------
def webpage_to_word(url, base_dir):
    html = fetch_html(url)
    soup = BeautifulSoup(html, "html.parser")

    title = extract_title(soup)

    folder, images_folder = create_page_folder(title, base_dir)
    filename = os.path.join(folder, sanitize_filename(title) + ".docx")

    content = extract_main_content(soup)

    image_map = download_images(content, url, images_folder)

    # Build document
    build_document(content, title, url, filename, image_map)

    print(f"✅ Saved: {filename}")
    return folder, title


def zip_directories(directories, output_zip_path):
    with zipfile.ZipFile(output_zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
        for directory in directories:
            root_name = os.path.basename(directory)
            for root, _, files in os.walk(directory):
                for filename in files:
                    file_path = os.path.join(root, filename)
                    archive_name = os.path.join(root_name, os.path.relpath(file_path, directory))
                    zipf.write(file_path, archive_name)


def create_docx_zip_from_urls(urls):
    if not urls:
        raise ValueError("At least one URL must be provided.")

    with tempfile.TemporaryDirectory() as temp_dir:
        created_paths = []
        titles = []

        for url in urls:
            folder_path, title = webpage_to_word(url, temp_dir)
            created_paths.append(folder_path)
            titles.append(title)

        zip_name = sanitize_filename(titles[0]) + ".zip" if len(created_paths) == 1 else "documents.zip"
        zip_path = os.path.join(temp_dir, zip_name)
        zip_directories(created_paths, zip_path)

        with open(zip_path, "rb") as f:
            data = f.read()

    return data, zip_name


# -----------------------------
# RUN ONE PAGE
# -----------------------------
def main_func(url):
    try:
        webpage_to_word(url, os.getcwd())
    except Exception as e:
        print(f"❌ Error processing {url}: {e}")
    time.sleep(5)